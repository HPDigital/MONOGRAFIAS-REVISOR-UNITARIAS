"""
MONOGRAFIAS REVISOR UNITARIAS
"""

#!/usr/bin/env python
# coding: utf-8

# In[ ]:


import os
from docx import Document
from openai import OpenAI
from dotenv import load_dotenv
from datetime import datetime

# 1️⃣ Cargar API Key
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# 2️⃣ Leer texto de DOCX
def leer_docx(ruta):
    doc = Document(ruta)
    return "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

# 3️⃣ Prompts de los agentes
prompts_agentes = {
    "TITULO": """Actúa como un revisor académico experto en metodología de investigación cuantitativa, no porpones ejemplos de mejoras o sujerencias.
Te limitas unicamente a hacer observaciones, no propones mejoras ni haces sujerencias.
Evalúa en el documento el TÍTULO de la monografía considerando:
- Claridad y precisión en el tema central.
- Identificación de variables independientes y dependientes.
- Delimitación geográfica y temporal.
- Coherencia con el objetivo general y el planteamiento del problema.
- Concisión (15-25 palabras).
- Ejemplo de estructura correcta: “Relación entre la frecuencia de compra y la rotación de inventario en Totto Potosí, gestión 2024”.
Entrega observaciones detalladas.""",

    "INTRODUCCIÓN": """Evalúa en el documento la INTRODUCCIÓN considerando el TÍTULO:
    Te limitas unicamente a hacer observaciones, no propones mejoras ni haces sujerencias.
- Presenta claramente el tema central.
- Explica el contexto general y la importancia del estudio.
- Introduce las variables, ámbito geográfico y temporal.
- Conecta con el problema de investigación.
- Justificación inicial coherente.
- Relación con CONTEXTO, OBJETIVOS y PLANTEAMIENTO DEL PROBLEMA.
Entrega observaciones detalladas.""",

    "CONTEXTO Y JUSTIFICACIÓN": """Evalúa en el documento CONTEXTO Y JUSTIFICACIÓN considerando TÍTULO e INTRODUCCIÓN:
    Te limitas unicamente a hacer observaciones, no propones mejoras ni haces sujerencias.
- Descripción clara de la empresa, sector o mercado.
- Relevancia para la gestión comercial y uso de Business Intelligence.
- Justificación Social donde se explica el aporte de esta monografia a la sociedad en genral.
- Justificación económica donde se explica el aporte en terminos de mejora econimica de esta monografia a la sociedad en genral
- Justificacion comercial comercial donde se explica el aporte de mejora comercial esta monografia a la empresa y la sociedad en genral.
- Inclusión de referencias o datos de apoyo.
- Coherencia con INTRODUCCIÓN, PLANTEAMIENTO DEL PROBLEMA y OBJETIVOS.
Entrega observaciones detalladas.""",

    "PLANTEAMIENTO DEL PROBLEMA": """Evalúa en el documento PLANTEAMIENTO DEL PROBLEMA considerando TÍTULO, INTRODUCCIÓN y CONTEXTO:
    Te limitas unicamente a hacer observaciones, no propones mejoras ni haces sujerencias.
- Definición clara y concreta del problema.
- Redacción como pregunta de investigación.
- Identificación de causas y consecuencias.
- Variables independientes y dependientes.
- Delimitación temporal y geográfica.
- Coherencia con OBJETIVOS y DIAGNÓSTICO.
- Se formula correctamente la pregunta problematica
Entrega observaciones detalladas.""",

    "OBJETIVOS": """Evalúa en el documento OBJETIVO GENERAL y OBJETIVOS ESPECÍFICOS:
    Te limitas unicamente a hacer observaciones, no propones mejoras ni haces sujerencias.
- Claridad y redacción con verbos en infinitivo.
- Relación directa con variables y problema.
- Delimitación temporal y geográfica.
- Secuencia lógica de objetivos específicos para alcanzar el general.
- Coherencia con PLANTEAMIENTO DEL PROBLEMA, ALCANCE y METODOLOGÍA.
Entrega observaciones detalladas.""",

    "ALCANCE Y LÍMITES": """Evalúa en el documento ALCANCE Y LÍMITES considerando OBJETIVOS y PROBLEMA:
    Te limitas unicamente a hacer observaciones, no propones mejoras ni haces sujerencias.
- Definición de población, variables, lugar y tiempo.
- Coherencia con objetivos y metodología.
- Diferenciación clara entre alcance (lo que cubre) y límites (lo que excluye).
- Razonabilidad de los límites definidos.
Entrega observaciones detalladas.""",

    "MARCO TEÓRICO": """Evalúa en el documento MARCO TEÓRICO REFERENCIAL considerando CONTEXTO, PROBLEMA y OBJETIVOS:
    Te limitas unicamente a hacer observaciones, no propones mejoras ni haces sujerencias.
- 6.1. Business Intelligence y Análisis de Datos: Conceptos clave, autores y su relación con el problema.
- 6.2. Funcionalidades de PowerBi para Análisis: Profundidad técnica y ejemplos aplicados.
- 6.3. Indicadores Clave de Desempeño (KPIs): Definición y explicacón de cada uno de los KPIs utilizados en el analisis y los graficos.
- 6.4. Modelos de Pronóstico en Excel: Explicación teórica y vínculo con objetivos.
- Uso de citas y formato APA.
- Coherencia general con el trabajo.
Entrega observaciones detalladas para cada subpunto.""",

    "DIAGNÓSTICO": """Evalúa en el documento DIAGNÓSTICO DE LA SITUACIÓN ACTUAL considerando CONTEXTO, PROBLEMA y OBJETIVOS:
    Te limitas unicamente a hacer observaciones, no propones mejoras ni haces sujerencias.
- 7.1. Descripción de la Empresa y Proceso de Ventas: Datos reales, tablas y gráficos.
- 7.2. Diccionario de datos: Explicacion y descripción clara de cada una de las variables numericas y categoricas con las que se hace el analisis de datos.
- 7.3. Estructura y Limpieza de la Base de Datos: Descripción clara de la preparación de datos.
- 7.4. Calidad y Consistencia: Evaluación de la integridad de los datos.
- Relación del diagnóstico con el problema.
- Coherencia con la METODOLOGÍA y ANÁLISIS DE DATOS.
Entrega observaciones detalladas para cada subpunto.""",

    "METODOLOGÍA": """Evalúa en el documento la METODOLOGÍA considerando OBJETIVOS y ALCANCE:
    Te limitas unicamente a hacer observaciones, no propones mejoras ni haces sujerencias.
- 8.1. Extracción y limpieza de Datos con Power Query: Procedimientos claros y replicables.
- 8.2. Modelado en PowerBi: Estructura de datos, relaciones y presentación de gráfico de relaciones.
- Claridad en pasos, herramientas utilizadas y replicabilidad.
- Coherencia con el DIAGNÓSTICO y ANÁLISIS DE DATOS.
Entrega observaciones detalladas para cada subpunto.""",

    "ANÁLISIS DE DATOS": """Evalúa en el documento el ANÁLISIS DE DATOS considerando METODOLOGÍA y OBJETIVOS:
    Te limitas unicamente a hacer observaciones, no propones mejoras ni haces sujerencias.
- 9.1. Análisis Descriptivo: Explicacion clara de los cuatro puntos siguientes a ser abordados.
- 9.2. Análisis Univariado: Presentación con gráficos, interpretación y explicación clara de estadísticas básica, interpretación y explicación clara de variables individuales numéricas y categóricas.
- 9.3. Análisis Bivariado y Correlaciones: Presentación con gráficos, interpretación y explicación clara de relaciones entre variables.
- 9.4. Análisis de Indicadores Comerciales: Presentación con gráficos, interpretación y explicación clara de la relacion enre las diferentes variables y la relevancia para la gestión.
- 9.5. Segmentación Avanzada: Uso de técnicas de agrupación y segmentación.
- Conexión de los resultados con los OBJETIVOS y PROBLEMA.
Entrega observaciones detalladas para cada subpunto.""",

    "MODELOS Y PRONÓSTICOS": """Evalúa en el documento GENERACIÓN DE MODELOS Y PRONÓSTICOS considerando ANÁLISIS DE DATOS y OBJETIVOS:
    Te limitas unicamente a hacer observaciones, no propones mejoras ni haces sujerencias.
- 10.1. Regresión Lineal con indicadores comerciales: Correcta aplicación y explicación.
- 10.2. Pronóstico con Función TENDENCIA: Pertinencia y resultados.
- Interpretación de resultados y conexión con los OBJETIVOS.
Entrega observaciones detalladas para cada subpunto.""",

    "CONCLUSIONES": """Evalúa en el documento CONCLUSIONES Y RECOMENDACIONES considerando todo el trabajo:
    Te limitas unicamente a hacer observaciones, no propones mejoras ni haces sujerencias.
- Responden al objetivo general y específicos.
- Resumen claro de los hallazgos principales.
- Recomendaciones prácticas y aplicables.
- Coherencia global con TÍTULO, PROBLEMA, ANÁLISIS y MODELOS.
Entrega observaciones detalladas.""",

}


# 4️⃣ Función de agente encadenado
def ejecutar_agente(seccion, prompt, texto, contexto_previos):
    contenido = prompt + "\n\n"
    if contexto_previos:
        contenido += "ANÁLISIS DE AGENTES PREVIOS:\n" + "\n".join(contexto_previos) + "\n\n"
    contenido += "TEXTO DE LA MONOGRAFÍA:\n" + texto

    respuesta = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": f"Eres un revisor experto en {seccion}."},
            {"role": "user", "content": contenido}
        ],
        temperature=0.2,
        max_tokens=5000
    )
    return respuesta.choices[0].message.content

# 5️⃣ Analizar monografía con contexto acumulado
def analizar_monografia_agentes(ruta_docx, carpeta_salida):
    texto = leer_docx(ruta_docx)
    nombre = os.path.splitext(os.path.basename(ruta_docx))[0]
    ruta_out = os.path.join(carpeta_salida, f"INFORME_REVISION_{nombre}.docx")

    doc = Document()
    doc.add_heading('Informe de Revisión de Monografía', 0)

    contexto_previos = []  # aquí se guarda el resultado de cada agente

    for seccion, prompt in prompts_agentes.items():
        doc.add_heading(seccion, level=1)
        resultado = ejecutar_agente(seccion, prompt, texto, contexto_previos)
        contexto_previos.append(f"[{seccion}]\n{resultado}")
        for linea in resultado.split("\n"):
            if linea.strip():
                doc.add_paragraph(linea.strip())

    doc.save(ruta_out)
    print(f"✅ Informe generado: {ruta_out}")
    return ruta_out

# 6️⃣ Procesar carpeta completa
def procesar_carpeta(carpeta_entrada, carpeta_salida):
    if not os.path.exists(carpeta_salida):
        os.makedirs(carpeta_salida)

    archivos = [f for f in os.listdir(carpeta_entrada) if f.endswith(".docx")]
    print(f"📂 Encontradas {len(archivos)} monografías.")

    for archivo in archivos:
        ruta = os.path.join(carpeta_entrada, archivo)
        analizar_monografia_agentes(ruta, carpeta_salida)

# 7️⃣ Ejecutar
if __name__ == "__main__":
    carpeta_monografias = r"C:\Users\HP\Downloads\MONOGRAFIAS\BENJAMIN MARTINEZ MARTINEZ"
    carpeta_informes = r"C:\Users\HP\Downloads\MONOGRAFIAS\BENJAMIN MARTINEZ MARTINEZ"
    procesar_carpeta(carpeta_monografias, carpeta_informes)

